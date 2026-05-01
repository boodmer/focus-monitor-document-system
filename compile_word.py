import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter
import config

IMG_DIR = config.IMG_DIR
DOC_DIR = config.DOC_DIR
CONTENT_DIR = config.CONTENT_DIR
OUT_DIR = config.OUT_DIR

# === OMML Equation Support ===
# Load Office XSLT for MathML -> OMML conversion (one-time)
_XSLT_PATH = r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL'
_xslt_tree = etree.parse(_XSLT_PATH)
_xslt_transform = etree.XSLT(_xslt_tree)

def latex_to_omml(latex_str):
    """Convert LaTeX math string to Word OMML XML element (native equation)."""
    # Step 1: LaTeX -> MathML
    mathml_str = latex2mathml.converter.convert(latex_str)
    # Step 2: MathML -> OMML via Office XSLT
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = _xslt_transform(mathml_tree)
    return omml_tree.getroot()

def add_omml_to_paragraph(paragraph, latex_str):
    """Insert a native Word equation into an existing paragraph (inline)."""
    omml_elem = latex_to_omml(latex_str)
    paragraph._element.append(omml_elem)

def add_display_equation(doc, latex_str):
    """Add a centered display equation paragraph (like ALT+= block equation)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    # Wrap in oMathPara for display-mode centering
    omml_elem = latex_to_omml(latex_str)
    omath_para = OxmlElement('m:oMathPara')
    omath_para.append(omml_elem)
    p._element.append(omath_para)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def set_update_fields_false(doc):
    """Set w:updateFields to false to prevent Word popup on open."""
    try:
        settings = doc.settings.element
        for uf in settings.findall(qn('w:updateFields')):
            settings.remove(uf)
        uf_elem = OxmlElement('w:updateFields')
        uf_elem.set(qn('w:val'), 'false')
        settings.append(uf_elem)
    except Exception as e:
        print(f"  Warning: Could not set updateFields: {e}")

def update_fields_via_word(filepath):
    """Use Word COM automation to update TOC, page numbers, and all fields."""
    try:
        import win32com.client
        abs_path = os.path.abspath(filepath)

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone

        wdoc = word.Documents.Open(abs_path)

        # Update all tables of contents
        for i in range(1, wdoc.TablesOfContents.Count + 1):
            wdoc.TablesOfContents.Item(i).Update()

        # Update all fields in main body
        wdoc.Content.Fields.Update()

        # Update fields in headers and footers
        for section in wdoc.Sections:
            for hf_idx in range(1, 4):
                try:
                    section.Headers(hf_idx).Range.Fields.Update()
                except Exception:
                    pass
                try:
                    section.Footers(hf_idx).Range.Fields.Update()
                except Exception:
                    pass

        wdoc.Save()
        wdoc.Close(False)
        word.Quit()

        print("  [OK] TOC and page numbers updated via Word automation")
        return True
    except ImportError:
        print("  [WARN] pywin32 not installed. Run: pip install pywin32")
        print("    To update TOC manually: Open in Word -> Ctrl+A -> F9")
        return False
    except Exception as e:
        print(f"  [WARN] Word automation failed: {e}")
        print("    To update TOC manually: Open in Word -> Ctrl+A -> F9")
        return False

def setup(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.3
    for s in doc.sections:
        s.page_width = Mm(210)
        s.page_height = Mm(297)
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(3.5)
        s.right_margin = Cm(2.0)
        
        # Add footer with page number (not on first page)
        s.different_first_page_header_footer = True
        footer = s.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        add_page_number(run)

    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(0)
        h_style.paragraph_format.space_after = Pt(0)
        h_style.paragraph_format.line_spacing = 1.3
        if i == 1:
            h_style.font.size = Pt(config.STYLES['chapter_title_size'])
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i == 2:
            h_style.font.size = Pt(config.STYLES['heading2_size'])
        else:
            h_style.font.size = Pt(config.STYLES['heading3_size'])
            h_style.font.italic = True

def heading(doc, text, level=1):
    if level == 1:
        text = text.upper()
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(0)
    h.paragraph_format.line_spacing = 1.3

    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.bold = True
        if level == 1:
            r.font.size = Pt(config.STYLES['chapter_title_size'])
        elif level == 2:
            r.font.size = Pt(config.STYLES['heading2_size'])
        else:
            r.font.size = Pt(config.STYLES['heading3_size'])

def para(doc, text):
    if not text.strip(): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('$') and part.endswith('$'):
            latex_str = part[1:-1]
            try:
                add_omml_to_paragraph(p, latex_str)
            except Exception as e:
                # Fallback: plain text if OMML conversion fails
                r = p.add_run(latex_str)
                r.font.italic = True
                r.font.name = 'Cambria Math'
                print(f"  [WARN] OMML fallback for: {latex_str}: {e}")
        else:
            if part:
                p.add_run(part)

def code_block(doc, lines):
    """Render a list of lines as a styled code block (Courier New, gray background)."""
    for code_line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        # Gray background via paragraph shading
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        # Code text in Courier New
        r = p.add_run(code_line if code_line else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1e, 0x1e, 0x2e)

def strip_number_prefix(title):
    """Remove leading number prefix like '2.7.1. ' from title."""
    import re
    title = re.sub(r'^Chương\s+\d+[\.\s]*', '', title)
    title = re.sub(r'^\d+(\.\d+)*\.?\s*', '', title)
    return title.strip()

# Global counters for auto-numbering
_ch_num = 0
_sec_num = 0
_subsec_num = 0

def parse_file(doc, filepath):
    global _ch_num, _sec_num, _subsec_num
    if not os.path.exists(filepath):
        print(f"  SKIP: {filepath}")
        return
    print(f"  Processing: {filepath}")
    with open(filepath, 'r', encoding='utf-8-sig') as f:
        lines_all = f.readlines()

    in_code_block = False
    code_lines = []

    for line in lines_all:
            line = line.rstrip('\n\r')
            stripped = line.strip()

            # --- Code block fence detection ---
            if stripped.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                else:
                    in_code_block = False
                    code_block(doc, code_lines)
                    code_lines = []
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if not stripped: continue

            # Titles -> page break + H1
            for prefix in ['CH0_TITLE|','CH1_TITLE|','CH2_TITLE|','CH3_TITLE|','CH4_TITLE|','CH5_TITLE|']:
                if stripped.startswith(prefix):
                    raw_title = stripped.split('|',1)[1]
                    clean = strip_number_prefix(raw_title).upper()
                    
                    # Check if it should be numbered
                    non_numbered = ['LỜI CẢM ƠN', 'TÀI LIỆU THAM KHẢO', 'TÓM TẮT', 'ABSTRACT', 'PHỤ LỤC']
                    is_numbered = True
                    for nn in non_numbered:
                        if clean.startswith(nn):
                            is_numbered = False
                            break
                    
                    if is_numbered and not stripped.startswith('CH0_'):
                        _ch_num += 1
                        _sec_num = 0
                        _subsec_num = 0
                        display = f"CHƯƠNG {_ch_num}: {clean}"
                    else:
                        display = clean
                    
                    doc.add_page_break()
                    heading(doc, display, level=1)
                    break
            else:
                # H2
                for prefix in ['CH0_H2|','CH1_H2|','CH2_H2|','CH3_H2|','CH4_H2|','CH5_H2|']:
                    if stripped.startswith(prefix):
                        _subsec_num = 0
                        raw_title = stripped.split('|',1)[1]
                        clean = strip_number_prefix(raw_title)
                        if clean:
                            clean = clean[0].upper() + clean[1:]
                        
                        if prefix == 'CH0_H2|':
                            heading(doc, clean, level=2)
                        else:
                            _sec_num += 1
                            display = f"{_ch_num}.{_sec_num}. {clean}"
                            heading(doc, display, level=2)
                        break
                else:
                    # H3 - Formatted as bold text to "merge" as requested
                    for prefix in ['CH0_H3|','CH1_H3|','CH2_H3|','CH3_H3|','CH4_H3|','CH5_H3|']:
                        if stripped.startswith(prefix):
                            raw_title = stripped.split('|',1)[1]
                            clean = strip_number_prefix(raw_title)
                            if clean:
                                clean = clean[0].upper() + clean[1:]
                            if prefix == 'CH0_H3|':
                                heading(doc, clean, level=3)
                            else:
                                _subsec_num += 1
                                display = f"{_ch_num}.{_sec_num}.{_subsec_num}. {clean}"
                                heading(doc, display, level=3)
                            break
                    else:
                        # Images
                        for prefix in ['CH2_IMAGE|','CH3_IMAGE|','CH4_IMAGE|']:
                            if stripped.startswith(prefix):
                                parts = stripped.split('|')
                                img_path = os.path.join(IMG_DIR, parts[1])
                                if os.path.exists(img_path):
                                    p = doc.add_paragraph()
                                    p.paragraph_format.space_after = Pt(0)
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    r = p.add_run()
                                    r.add_picture(img_path, width=Inches(5.5))
                                    cap = doc.add_paragraph(parts[2], style='ImgCaption')
                                    cap.paragraph_format.space_after = Pt(0)
                                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                break
                        else:
                            # Handle Equations — native Word OMML
                            for prefix in ['CH2_EQ|','CH3_EQ|','CH4_EQ|']:
                                if stripped.startswith(prefix):
                                    math_latex = stripped.split('|', 1)[1]
                                    try:
                                        add_display_equation(doc, math_latex)
                                    except Exception as e:
                                        print(f"  [WARN] Equation OMML error: {e}")
                                        para(doc, f"[Equation: {math_latex}]")
                                    break
                            else:
                                # Screenshot placeholders
                                if stripped.startswith('[SCREENSHOT]'):
                                    caption = stripped.replace('[SCREENSHOT]', '').strip()
                                    # Add empty bordered paragraph as placeholder
                                    p = doc.add_paragraph()
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    r = p.add_run('\n\n\n[Chèn ảnh chụp màn hình tại đây]\n\n\n')
                                    r.font.size = Pt(12)
                                    r.font.italic = True
                                    r.font.color.rgb = None
                                    # Add border to paragraph
                                    from docx.oxml.ns import qn
                                    from docx.oxml import OxmlElement
                                    pPr = p._element.get_or_add_pPr()
                                    pBdr = OxmlElement('w:pBdr')
                                    for border_name in ['top', 'left', 'bottom', 'right']:
                                        border = OxmlElement(f'w:{border_name}')
                                        border.set(qn('w:val'), 'single')
                                        border.set(qn('w:sz'), '4')
                                        border.set(qn('w:space'), '4')
                                        border.set(qn('w:color'), '999999')
                                        pBdr.append(border)
                                    pPr.append(pBdr)
                                    # Caption
                                    cap = doc.add_paragraph(caption, style='ImgCaption')
                                    cap.paragraph_format.space_after = Pt(0)
                                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    para(doc, stripped)

def generate_cover_page(doc):
    """Hàm tạo trang bìa chuẩn cho Word (Anh có thể chỉnh sửa khoảng cách, text ở đây)"""

    # Nội dung trang bìa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(f"{config.META['department']}\n{config.META['university']}")
    r.bold = True; r.font.size = Pt(14)
    
    doc.add_paragraph("─────────────────────").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(2): doc.add_paragraph("")
    
    p = doc.add_paragraph(config.META['degree'].upper())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.bold = True; r.font.size = Pt(18)
    
    p = doc.add_paragraph(f"Ngành: {config.META['major']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(14)
    
    for _ in range(2): doc.add_paragraph("")
    p = doc.add_paragraph(config.META['title'].upper())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.bold = True; r.font.size = Pt(18)
    
    for _ in range(3): doc.add_paragraph("")
    
    # Use table for exact alignment of info block
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Set column widths manually
    for row in table.rows:
        row.cells[0].width = Inches(1.0) # Empty indent column
        row.cells[1].width = Inches(2.2) # Label column
        row.cells[2].width = Inches(3.0) # Value column
    
    def set_cell(row_idx, label, value):
        cell_label = table.cell(row_idx, 1)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(14)
        
        cell_val = table.cell(row_idx, 2)
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(value)
        r.font.size = Pt(14)

    set_cell(0, "Giảng viên hướng dẫn:", "TS. Phạm Huy Thông")
    set_cell(1, "Sinh viên thực hiện:", config.META['student_name'])
    set_cell(2, "Mã số sinh viên:", config.META['student_id'])
    set_cell(3, "Lớp:", config.META['class_name'])
    
    for _ in range(2): doc.add_paragraph("")
    
    p = doc.add_paragraph(f"{config.META['location']} - {config.META['year']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.bold = True; r.font.size = Pt(14)
    
    doc.add_page_break()

def build():
    doc = Document()
    
    # Create ImgCaption style
    styles = doc.styles
    if 'ImgCaption' not in styles:
        style = styles.add_style('ImgCaption', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(13)
        style.font.italic = True
        style.font.color.rgb = RGBColor(0,0,0)

    try:
        for i in range(1, 4):
            if f'TOC {i}' in doc.styles:
                doc.styles[f'TOC {i}'].font.size = Pt(13)
                doc.styles[f'TOC {i}'].font.name = 'Times New Roman'
                doc.styles[f'TOC {i}'].font.italic = False
                doc.styles[f'TOC {i}'].font.bold = False
    except:
        pass

    setup(doc)

    # === BÌA ===
    generate_cover_page(doc)

    # === PHIẾU GIAO ĐỀ TÀI ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n\n")
    r.bold = True
    r2 = p.add_run("PHIẾU GIAO ĐỀ TÀI ĐỒ ÁN TỐT NGHIỆP")
    r2.bold = True; r2.font.size = Pt(16)
    doc.add_paragraph("")
    para(doc, f"Họ tên sinh viên: {config.META['student_name']}          Mã SV: {config.META['student_id']}")
    para(doc, f"Lớp: {config.META['class_name']}          Ngành: {config.META['major']}          Khóa: {config.META['cohort']}")
    p = doc.add_paragraph()
    p.add_run("Tên đề tài: ").bold = True
    p.add_run(config.META['title'])
    p = doc.add_paragraph()
    p.add_run("Mục tiêu đề tài: ").bold = True
    p.add_run(config.META['goal'])
    para(doc, f"Thời gian thực hiện: {config.META['duration']}")
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph("NGƯỜI HƯỚNG DẪN                              TRƯỞNG ĐƠN VỊ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.bold = True
    doc.add_page_break()

    # === MỤC LỤC ===
    p = doc.add_paragraph("MỤC LỤC")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.bold = True; r.font.size = Pt(config.STYLES['chapter_title_size'])
    
    
    # Add dynamic TOC field
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    doc.add_page_break()

    # === DANH SÁCH HÌNH VẼ ===
    heading(doc, "DANH SÁCH HÌNH VẼ", level=1)
    
    
    # Add dynamic LOF field
    lof_para = doc.add_paragraph()
    run = lof_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\t "ImgCaption,1"'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    doc.add_page_break()

    # === DANH MỤC VIẾT TẮT ===
    heading(doc, "DANH MỤC CÁC TỪ VIẾT TẮT", level=1)
    abbrevs = config.ABBREVIATIONS
    table = doc.add_table(rows=len(abbrevs)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, t in enumerate(['STT', 'Viết tắt', 'Ý nghĩa']):
        hdr[i].text = t
        for r in hdr[i].paragraphs[0].runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    for idx, (abbr, meaning) in enumerate(abbrevs):
        row = table.rows[idx+1].cells
        row[0].text = str(idx+1)
        row[1].text = abbr
        row[2].text = meaning
        for cell in row:
            for r in cell.paragraphs[0].runs: r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    doc.add_page_break()

    # === DANH MỤC KÝ HIỆU ===
    heading(doc, "DANH MỤC CÁC KÝ HIỆU", level=1)
    symbols = config.SYMBOLS_WORD
    table2 = doc.add_table(rows=len(symbols)+1, cols=3)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    for i, t in enumerate(['STT', 'Ký hiệu', 'Ý nghĩa']):
        hdr2[i].text = t
        for r in hdr2[i].paragraphs[0].runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    for idx, (sym, meaning) in enumerate(symbols):
        row = table2.rows[idx+1].cells
        row[0].text = str(idx+1)
        row[1].text = sym
        row[2].text = meaning
        for cell in row:
            for r in cell.paragraphs[0].runs: r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    # === NỘI DUNG CHÍNH ===
    chapter_files = [os.path.join(CONTENT_DIR, f) for f in config.CHAPTER_FILES]

    for cf in chapter_files:
        parse_file(doc, cf)

    out = os.path.join(OUT_DIR, '2024600002_NguyenQuangMinh_DATN.docx')
    
    doc.save(out)
    print(f"\n=== SAVED: {out} ===")
    
    # Use Word COM automation to update TOC/LOF with correct page numbers
    print("  Updating fields via Word automation...")
    if update_fields_via_word(out):
        # COM succeeded: fields are populated, suppress future update prompts
        doc_reopen = Document(out)
        set_update_fields_false(doc_reopen)
        doc_reopen.save(out)
        print("  Popup suppressed (fields already up-to-date)")
    else:
        print("  Word will prompt to update fields on open (Ctrl+A -> F9)")
    print(f"\n=== DONE ===")

if __name__ == '__main__':
    build()
